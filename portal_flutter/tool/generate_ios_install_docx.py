#!/usr/bin/env python3
"""Генерирует пошаговую инструкцию по установке iOS-клиента в Word (.docx)."""
from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
except ImportError as e:
    raise SystemExit("Установи: pip install python-docx") from e

OUT = Path(__file__).resolve().parent.parent / "docs" / "Portal_iOS_установка.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step).font.size = Pt(11)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    t = doc.add_paragraph()
    r = t.add_run("Portal — установка на iPhone / iPad")
    r.bold = True
    r.font.size = Pt(16)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    add_heading(doc, "Важно до начала", level=1)
    add_para(
        doc,
        "На iOS нельзя скачать один файл «как APK на Android». Apple не раздаёт приложения таким способом. "
        "Нужен Mac с Xcode или установка через TestFlight/App Store (отдельная история).",
    )
    add_para(
        doc,
        "На Android проще: скачиваешь Portal-Flutter.apk из релиза GitHub и открываешь на телефоне.",
    )

    add_heading(doc, "Что понадобится (способ через Mac)", level=1)
    add_numbered_steps(
        doc,
        [
            "Компьютер Mac с установленным Xcode (из App Store на Mac).",
            "Установленный Flutter SDK (flutter.dev) и чтобы команда flutter работала в Терминале.",
            "Кабель USB для iPhone (или доверенная сеть для беспроводной отладки — позже, сначала проще кабель).",
            "Apple ID (можно бесплатный; для личной установки на свой телефон достаточно).",
        ],
    )

    add_heading(doc, "Способ 1 — собрать и поставить через Xcode (рекомендуется)", level=1)

    add_heading(doc, "Часть A. Подготовка проекта на Mac", level=2)
    add_numbered_steps(
        doc,
        [
            "Открой Safari или другой браузер и зайди на страницу репозитория Portal на GitHub (тот, куда залит код).",
            "Нажми зелёную кнопку «Code» → скопируй HTTPS-ссылку для git clone.",
            "Открой «Терминал» (Spotlight: Cmd+Пробел, набери Terminal, Enter).",
            "Перейди в папку, куда хочешь положить проект, например: cd ~/Desktop",
            "Выполни: git clone <вставь скопированную ссылку>",
            "Перейди в папку репозитория: cd portal (или как называется папка после clone).",
            "Перейди в мобильное приложение: cd portal_flutter",
            "Сгенерируй iOS-платформу (один раз или после чистки): "
            "flutter create . --project-name portal_flutter --org org.portal --platforms=ios",
            "Примени патч для iOS: python3 tool/patch_ios_info_plist.py",
            "Подтяни зависимости: flutter pub get",
        ],
    )

    add_heading(doc, "Часть B. Xcode: подпись и запуск на телефон", level=2)
    add_numbered_steps(
        doc,
        [
            "В Терминале (всё ещё в папке portal_flutter) выполни: open ios/Runner.xcworkspace",
            "Откроется Xcode. В левой колонке выбери проект Runner (синяя иконка вверху).",
            "В центре открой вкладку «Signing & Capabilities».",
            "Поставь галочку «Automatically manage signing», если снята.",
            "В поле «Team» выбери свою команду (Add an Account… — войди под Apple ID, если пусто).",
            "Сверху по центру Xcode: в списке схем выбери «Runner», а справа от него — свой iPhone (подключи кабелем и разблокируй телефон).",
            "На iPhone при появлении запроса «Доверять этому компьютеру?» нажми «Доверять».",
            "В Xcode нажми треугольник «Run» (▶) или Cmd+R — начнётся сборка и установка на телефон.",
            "Если iOS пишет, что разработчик не доверен: Настройки → Основные → VPN и управление устройством "
            "(или «Управление устройством») → твой Apple ID → Доверять.",
        ],
    )

    add_para(
        doc,
        "С бесплатным Apple ID подпись обычно действует около 7 дней — потом нужно снова собрать/установить из Xcode.",
        bold=False,
    )

    add_heading(doc, "Способ 2 — архив из GitHub Actions (Portal-Flutter-iOS-nosign.zip)", level=1)
    add_para(
        doc,
        "В Actions → workflow «Portal Flutter Build» → job ios-nosign публикуется архив Portal-Flutter-iOS-nosign.zip. "
        "Внутри лежит Runner.app — это не «установщик в один клик».",
    )
    add_numbered_steps(
        doc,
        [
            "Зайди на GitHub → вкладка «Actions» → выбери успешный запуск «Portal Flutter Build».",
            "Внизу страницы запуска в блоке «Artifacts» скачай Portal-Flutter-iOS-nosign.zip.",
            "Распакуй zip на Mac — появится папка Runner.app.",
            "Чтобы поставить на телефон, всё равно нужна подпись: проще открыть исходники и собрать по Способу 1. "
            "Артефакт чаще нужен для отладки или если ты уже умеешь подписывать через Xcode/Organizer.",
        ],
    )

    add_heading(doc, "Если Mac нет", level=1)
    add_para(
        doc,
        "Установить «как APK» без Mac нельзя. Варианты: попросить знакомого со Mac собрать по инструкции выше; "
        "или использовать Android-версию для приёма в фоне; или в будущем — TestFlight/App Store при публикации.",
    )

    add_heading(doc, "Сеть и пароль", level=1)
    add_numbered_steps(
        doc,
        [
            "ПК и iPhone должны видеть друг друга (одна Wi‑Fi или Tailscale).",
            "На ПК запущен Portal, порт 12345; в приложении на iPhone тот же пароль сети, что в настройках Portal на ПК.",
            "На iPhone во вкладке «Пиры» при необходимости укажи IP ПК (для Tailscale часто 100.x.x.x).",
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        f"Файл сгенерирован скриптом tool/generate_ios_install_docx.py. Актуальные детали также в IOS_INSTALL.md.",
        bold=False,
    )

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
